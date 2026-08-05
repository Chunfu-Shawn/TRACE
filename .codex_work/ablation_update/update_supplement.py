from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE = Path(
    "/Users/chunfu/Desktop/BGM_lab/translation_model/TRACE/.codex_work/ablation_update/"
    "supplementary_information.docx"
)
OUTPUT = SOURCE.with_name("supplementary_information.updated.docx")


REPLACEMENTS = {
    "Structural ablations retained TRACE's nucleotide embedding": (
        "All four models used the same nucleotide input, prediction head, losses, "
        "chromosome partitions and optimization settings. TRACE-Real and TRACE shared a "
        "12-layer AdaLN-Zero Transformer (model dimension, 384; attention heads, 16; "
        "feed-forward dimension, 768; dropout, 0.1). A 16,840-gene expression vector and "
        "species embedding were compressed through a 64-dimensional bottleneck to a "
        "16-dimensional cellular-context representation that modulated each block. LN "
        "replaced AdaLN-Zero with standard pre-LayerNorm self-attention at the same depth "
        "and width and omitted cellular-context and species inputs. Conv instead used 12 "
        "sequence-only residual one-dimensional convolutional blocks (kernel width, 7; "
        "hidden and feed-forward dimensions, 384; GELU; dropout, 0.1)."
    ),
    "The cellular-context ablation used the TRACE backbone.": (
        "All models were trained for up to 50 epochs on the same five human cellular "
        "contexts. TRACE-Real received unperturbed expression vectors. For TRACE, 15% of "
        "training samples received an exact zero vector; 30% of the remaining vectors were "
        "scaled by a value sampled uniformly from 0 to 1, after which Gaussian noise "
        "(s.d., 0.1) was added to every non-zero vector. LN and Conv received no "
        "cellular-context input. The objective was micro loss + 2.0 × macro loss + 0.2 × "
        "ranking loss, with the macro-loss weight increased from 0.2 to 2.0 during warm-up. "
        "The checkpoint with the lowest validation loss was selected for each model."
    ),
    "Each strategy was trained on matched sets of 5, 22 or 40 human cellular contexts": (
        "Validation-selected checkpoints were evaluated on chromosome-held-out test "
        "transcripts from 26 human cell lines absent from training. Profiles with RPF depth "
        "≥ 0.1 were retained, and context-level metrics were calculated for the 14 cell "
        "lines with at least 800 retained transcripts. Nucleotide-profile Spearman ρ, "
        "CDS-mean signal Spearman ρ and CDS-mean MAE were averaged equally across "
        "eligible contexts."
    ),
    "Supplementary Table 2 | Selected model-ablation performance": (
        "Supplementary Table 2 | Ablation performance on test data from 26 unseen cell lines"
    ),
    "Table note. Values are the best values recorded across 50 epochs": (
        "Table note. Values are means across the 14 eligible cellular contexts among the "
        "26 unseen human cell lines in the chromosome-held-out test dataset. Eligibility "
        "required at least 800 retained transcripts with RPF depth of at least 0.1 per "
        "context. Checkpoints were selected solely by the lowest validation loss before "
        "test evaluation (epochs 29, 41, 42 and 5 for TRACE-Real, TRACE, LN and Conv, "
        "respectively). Higher values are better for both Spearman ρ metrics, whereas "
        "lower values are better for CDS-mean MAE. The best value in each column is shaded "
        "green. TRACE corresponds to TRACE-Augment in the source result files."
    ),
}


def replace_matching_paragraph(document, prefix, text):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {prefix!r}, found {len(matches)}")
    paragraph = matches[0]
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def clear_cell_shading(cell):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is not None:
        properties.remove(shading)


def remove_last_column(table):
    for row in table.rows:
        row._tr.remove(row.cells[-1]._tc)
    grid = table._tbl.tblGrid
    grid.remove(grid.gridCol_lst[-1])


def set_cell_text(cell, text, bold=False):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


document = Document(SOURCE)
for prefix, replacement in REPLACEMENTS.items():
    replace_matching_paragraph(document, prefix, replacement)

table = document.tables[0]
if len(table.columns) != 5 or len(table.rows) != 5:
    raise RuntimeError("Supplementary Table 2 does not have the expected 5 x 5 structure")
remove_last_column(table)

headers = [
    "Model",
    "Nucleotide-profile\nSpearman ρ",
    "CDS-mean signal\nSpearman ρ",
    "CDS-mean MAE",
]
rows = [
    ("TRACE-Real", "0.4381", "0.3432", "0.1648"),
    ("TRACE", "0.4430", "0.3590", "0.1699"),
    ("LN", "0.4386", "0.3431", "0.1622"),
    ("Conv", "0.3604", "0.2911", "0.1698"),
]

for column, text in enumerate(headers):
    set_cell_text(table.cell(0, column), text, bold=True)
for row_index, values in enumerate(rows, start=1):
    for column, text in enumerate(values):
        set_cell_text(table.cell(row_index, column), text, bold=(column == 0))

for row in table.rows:
    for cell in row.cells:
        clear_cell_shading(cell)
set_cell_shading(table.cell(2, 1), "C6E0B4")
set_cell_shading(table.cell(2, 2), "C6E0B4")
set_cell_shading(table.cell(3, 3), "C6E0B4")

table.autofit = False
column_widths = [Inches(1.25), Inches(1.95), Inches(1.95), Inches(1.72)]
for row in table.rows:
    for column, width in enumerate(column_widths):
        row.cells[column].width = width

document.save(OUTPUT)
print(OUTPUT)
