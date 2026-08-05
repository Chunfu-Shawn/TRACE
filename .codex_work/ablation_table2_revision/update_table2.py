from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SOURCE = Path(__file__).with_name("supplementary_information.docx")
OUTPUT = SOURCE.with_name("supplementary_information.updated.docx")


def set_cell_text(cell, text, bold=False, font_size=7.2):
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clear_cell_shading(cell):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is not None:
        properties.remove(shading)


def set_cell_shading(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def remove_vertical_borders(cell):
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge_name in ("left", "right", "start", "end"):
        edge = borders.find(qn(f"w:{edge_name}"))
        if edge is None:
            edge = OxmlElement(f"w:{edge_name}")
            borders.append(edge)
        edge.set(qn("w:val"), "nil")


document = Document(SOURCE)
table = document.tables[0]
if len(table.rows) != 5 or len(table.columns) != 5:
    raise RuntimeError("Supplementary Table 2 does not have the expected 5 x 5 structure")

# Two validation metrics are added while the test MAE column is replaced by
# CDS-profile Spearman, yielding seven columns in total.
table.add_column(Inches(0.9))
table.add_column(Inches(0.9))

headers = [
    "Model",
    "Best validation\nloss (epoch)",
    "Validation mean RNA-profile\nSpearman ρ (epoch)",
    "Validation CDS-mean\nSpearman ρ (epoch)",
    "Test mean RNA-profile\nSpearman ρ",
    "Test mean CDS-profile\nSpearman ρ",
    "Test CDS-mean\nSpearman ρ",
]
rows = [
    ("TRACE-Real", "0.1067 (29)", "0.4459 (47)", "0.6487 (41)", "0.4381", "0.2908", "0.3432"),
    ("TRACE", "0.1077 (41)", "0.4441 (40)", "0.6345 (42)", "0.4430", "0.3095", "0.3590"),
    ("LN", "0.1109 (42)", "0.4270 (10)", "0.6001 (44)", "0.4386", "0.2922", "0.3431"),
    ("Conv", "0.1190 (5)", "0.3987 (5)", "0.5867 (40)", "0.3604", "0.3061", "0.2911"),
]

for column, header in enumerate(headers):
    set_cell_text(table.cell(0, column), header, bold=True, font_size=6.8)
for row_index, values in enumerate(rows, start=1):
    for column, value in enumerate(values):
        set_cell_text(table.cell(row_index, column), value, bold=(column == 0))

for row in table.rows:
    for cell in row.cells:
        clear_cell_shading(cell)
        remove_vertical_borders(cell)

green = "C6E0B4"
for column in (1, 2, 3):
    set_cell_shading(table.cell(1, column), green)
for column in (4, 5, 6):
    set_cell_shading(table.cell(2, column), green)

table.autofit = False
column_widths = [
    Inches(0.76),
    Inches(0.82),
    Inches(1.23),
    Inches(1.12),
    Inches(1.00),
    Inches(1.00),
    Inches(0.94),
]
for row in table.rows:
    for column, width in enumerate(column_widths):
        row.cells[column].width = width

title_matches = [
    paragraph
    for paragraph in document.paragraphs
    if paragraph.text.startswith("Supplementary Table 2 |")
]
if len(title_matches) != 1:
    raise RuntimeError(f"Expected one Supplementary Table 2 title, found {len(title_matches)}")
title = title_matches[0]
title_style = title.style
title.clear()
title.style = title_style
title.add_run(
    "Supplementary Table 2 | Ablation performance during validation and zero-shot evaluation"
)

notes = [
    paragraph
    for paragraph in document.paragraphs
    if paragraph.text.startswith("Table note. Best validation loss")
]
if len(notes) != 1:
    raise RuntimeError(f"Expected one Supplementary Table 2 note, found {len(notes)}")
note = notes[0]
note_style = note.style
note.clear()
note.style = note_style
note.add_run(
    "Table note. Validation values are the best values recorded independently for each "
    "metric across 50 training epochs on the human_5c_6k_depth0.1_cov0.1_rpm1 validation "
    "dataset; the corresponding epoch is shown in parentheses. Test values are equal-context "
    "means across the 14 eligible cellular contexts among 26 unseen human cell lines in the "
    "chromosome-held-out test dataset, using the checkpoint selected by the lowest validation "
    "loss. Eligibility required at least 800 retained transcripts with RPF depth of at least "
    "0.1 per context. Higher Spearman ρ and lower validation loss indicate better performance. "
    "The best value in each column is shaded green. TRACE-Real achieved the best validation "
    "values, whereas TRACE achieved the best zero-shot test values, consistent with expression "
    "masking and interpolation improving generalization to unseen cellular contexts. TRACE "
    "corresponds to TRACE-Augment in the source result files."
)

evaluation_paragraphs = [
    paragraph
    for paragraph in document.paragraphs
    if paragraph.text.startswith("Validation-selected checkpoints were evaluated")
]
if len(evaluation_paragraphs) != 1:
    raise RuntimeError(
        f"Expected one ablation evaluation paragraph, found {len(evaluation_paragraphs)}"
    )
paragraph = evaluation_paragraphs[0]
paragraph_style = paragraph.style
paragraph.clear()
paragraph.style = paragraph_style
paragraph.add_run(
    "Validation-selected checkpoints were evaluated on chromosome-held-out test transcripts "
    "from 26 human cell lines absent from training. Profiles with RPF depth ≥ 0.1 were retained. "
    "Mean RNA-profile Spearman ρ, mean CDS-profile Spearman ρ and CDS-mean Spearman ρ were "
    "calculated for the 14 cell lines with at least 800 retained transcripts and averaged equally "
    "across eligible contexts."
)

document.save(OUTPUT)
print(OUTPUT)
