import csv
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


WORKDIR = Path(__file__).resolve().parent
CSV_PATH = Path(
    "/Users/chunfu/Desktop/BGM_lab/translation_model/results/ablation/"
    "loss_curves/model_metric_comparison.summary.csv"
)
MAIN_SOURCE = WORKDIR / "TRACE_manuscript.source.docx"
SUPP_SOURCE = WORKDIR / "supplementary_information.source.docx"
MAIN_OUTPUT = WORKDIR / "TRACE_manuscript.revised.docx"
SUPP_OUTPUT = WORKDIR / "supplementary_information.revised.docx"


METRICS = [
    ("best_train_loss", "best_train_loss_epoch", "Best train loss\n(epoch)", 4, "min"),
    ("best_valid_loss", "best_valid_loss_epoch", "Best validation loss\n(epoch)", 4, "min"),
    ("best_profile_spearman", "best_profile_spearman_epoch", "Profile Spearman ρ\n(epoch)", 3, "max"),
    (
        "best_cds_mean_scale_spearman",
        "best_cds_mean_scale_spearman_epoch",
        "CDS-mean scale\nSpearman ρ (epoch)",
        3,
        "max",
    ),
    ("best_cds_mean_mae", "best_cds_mean_mae_epoch", "CDS-mean MAE\n(epoch)", 4, "min"),
]

DISPLAY_NAMES = {
    "TRACE Zero": "TRACE Zero",
    "TRACE Real": "TRACE Real",
    "TRACE Mask+Interp.": "TRACE Augment",
    "LN Transformer": "LN-model",
    "Conv model (5c)": "Conv-model",
}


def paragraph_starting_with(document, prefix):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {prefix!r}, found {len(matches)}")
    return matches[0]


def replace_paragraph_text(paragraph, text, bold=False):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = bold
    return run


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("left", start), ("bottom", bottom), ("right", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, top=None, bottom=None, left="nil", right="nil"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for side, value in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        if value is None:
            value = "nil"
        edge = borders.find(qn(f"w:{side}"))
        if edge is None:
            edge = OxmlElement(f"w:{side}")
            borders.append(edge)
        edge.set(qn("w:val"), value)
        if value == "single":
            edge.set(qn("w:sz"), "12")
            edge.set(qn("w:space"), "0")
            edge.set(qn("w:color"), "auto")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    justification = tbl_pr.find(qn("w:jc"))
    if justification is None:
        justification = OxmlElement("w:jc")
        tbl_pr.append(justification)
    justification.set(qn("w:val"), "center")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])


def add_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


with CSV_PATH.open(newline="", encoding="utf-8-sig") as handle:
    rows = list(csv.DictReader(handle))

if len(rows) != 5:
    raise RuntimeError(f"Expected five ablation rows, found {len(rows)}")
if len({row["validation_dataset"] for row in rows}) != 1:
    raise RuntimeError("Validation dataset is not constant across ablation rows")
if len({row["epochs_recorded"] for row in rows}) != 1:
    raise RuntimeError("Number of recorded epochs is not constant across ablation rows")

best_by_metric = {}
for value_key, _, _, _, direction in METRICS:
    values = [float(row[value_key]) for row in rows]
    best_by_metric[value_key] = min(values) if direction == "min" else max(values)

main = Document(MAIN_SOURCE)
supp = Document(SUPP_SOURCE)

# The prior-work table moves from Supplementary Table 2 to Supplementary Table 3.
for paragraph in main.paragraphs:
    if "Supplementary Table 2" in paragraph.text:
        replace_paragraph_text(
            paragraph,
            paragraph.text.replace("Supplementary Table 2", "Supplementary Table 3"),
        )

main_ablation_summary = paragraph_starting_with(main, "Cellular-context conditioning was further evaluated")
if "Supplementary Table 2" not in main_ablation_summary.text:
    replace_paragraph_text(
        main_ablation_summary,
        main_ablation_summary.text.replace(
            "Full model configurations, augmentation parameters and evaluation procedures are provided in Supplementary Methods.",
            "Results for the five-context model ablation are summarized in Supplementary Table 2. Full model configurations, augmentation parameters and evaluation procedures are provided in Supplementary Methods.",
        ),
    )

old_caption = paragraph_starting_with(supp, "Supplementary Table 2 | Capability comparison")
replace_paragraph_text(
    old_caption,
    old_caption.text.replace("Supplementary Table 2", "Supplementary Table 3"),
    bold=True,
)

# Add the new caption using the established supplementary-table caption formatting.
new_caption = supp.add_paragraph(style=old_caption.style)
new_caption_ppr = new_caption._p.get_or_add_pPr()
new_caption._p.remove(new_caption_ppr)
new_caption._p.insert(0, deepcopy(old_caption._p.pPr))
new_caption.paragraph_format.page_break_before = False
replace_paragraph_text(
    new_caption,
    "Supplementary Table 2 | Model-ablation performance on the five-context dataset",
    bold=True,
)

table = supp.add_table(rows=len(rows) + 1, cols=len(METRICS) + 1)
table.style = "Table Grid"
widths = [2492, 2500, 2700, 2500, 2800, 2500]
set_table_geometry(table, widths)
add_repeat_table_header(table.rows[0])

headers = ["Model"] + [metric[2] for metric in METRICS]
for column_index, header in enumerate(headers):
    cell = table.cell(0, column_index)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell, top=85, start=70, bottom=85, end=70)
    set_cell_borders(cell, top="single", bottom="single")
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(header)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)

for row_index, source_row in enumerate(rows, start=1):
    model_cell = table.cell(row_index, 0)
    model_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(model_cell)
    set_cell_borders(model_cell, bottom="single" if row_index == len(rows) else None)
    model_paragraph = model_cell.paragraphs[0]
    model_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    model_paragraph.paragraph_format.space_before = Pt(0)
    model_paragraph.paragraph_format.space_after = Pt(0)
    model_run = model_paragraph.add_run(DISPLAY_NAMES[source_row["model"]])
    model_run.bold = True
    model_run.font.size = Pt(10)

    for metric_index, (value_key, epoch_key, _, precision, _) in enumerate(METRICS, start=1):
        cell = table.cell(row_index, metric_index)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        set_cell_borders(cell, bottom="single" if row_index == len(rows) else None)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        value = float(source_row[value_key])
        epoch = int(source_row[epoch_key])
        run = paragraph.add_run(f"{value:.{precision}f} ({epoch})")
        run.font.size = Pt(10)
        if abs(value - best_by_metric[value_key]) < 1e-12:
            run.bold = True
            run.font.color.rgb = RGBColor(27, 77, 27)
            set_cell_shading(cell, "C6E0B4")

note = supp.add_paragraph(style=old_caption.style)
note.paragraph_format.space_before = Pt(6)
note.paragraph_format.space_after = Pt(1)
note_text = (
    "Table note. Values are the best values recorded across 50 epochs, with the corresponding epoch in parentheses. "
    "Training loss was computed on the training set; all other metrics used the human_5c_6k_depth0.1_cov0.1_rpm1 validation set. "
    "Lower values are better for losses and CDS-mean MAE, whereas higher values are better for Spearman ρ. "
    "The best value in each column is shaded green. Because each metric was selected independently, values within a row can correspond to different checkpoints. "
    "Model names were standardized to the manuscript terminology: TRACE Augment corresponds to TRACE Mask+Interp. in the source file, LN-model to LN Transformer and Conv-model to Conv model (5c)."
)
note_run = note.add_run(note_text)
note_run.font.size = Pt(7.5)
note_run.font.color.rgb = RGBColor(0, 0, 0)

# Move the new caption, table and note immediately before Supplementary Table 3.
old_caption._p.addprevious(new_caption._p)
old_caption._p.addprevious(table._tbl)
old_caption._p.addprevious(note._p)

main.save(MAIN_OUTPUT)
supp.save(SUPP_OUTPUT)

# Structural verification.
main_check = Document(MAIN_OUTPUT)
supp_check = Document(SUPP_OUTPUT)
main_text = "\n".join(paragraph.text for paragraph in main_check.paragraphs)
supp_text = "\n".join(paragraph.text for paragraph in supp_check.paragraphs)
assert main_text.count("Supplementary Table 2") == 1
assert main_text.count("Supplementary Table 3") == 2
assert "Supplementary Table 2 | Model-ablation performance" in supp_text
assert "Supplementary Table 3 | Capability comparison" in supp_text
assert len(supp_check.tables) == 2
assert len(supp_check.tables[0].rows) == 6
assert len(supp_check.tables[0].columns) == 6
assert len(supp_check.tables[1].rows) == 12
assert len(supp_check.tables[1].columns) == 10

print(MAIN_OUTPUT)
print(SUPP_OUTPUT)
