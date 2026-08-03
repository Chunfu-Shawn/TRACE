from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


WORKDIR = Path(".codex_scope_sync")


OUTLINE_TEXT = {
    3: "Translation varies across transcripts and cellular contexts, so transcript abundance alone cannot predict protein output. Nucleotide-resolution occupancy models could support both sequence design and translated-ORF discovery.",
    4: "Ribo-seq provides nucleotide-resolution occupancy but is difficult to deploy broadly. Existing models usually separate TE estimation, local dynamics and ORF identification, and often require CDS annotations, fixed windows or Ribo-seq input.",
    5: "A unified model must capture multiscale sequence features and cellular context while learning from heterogeneous Ribo-seq data. We address this with harmonized targets, variable-length modeling, context conditioning and chromosome-held-out evaluation.",
    6: "TRACE predicts full-length occupancy profiles and is annotation-free at inference. Trained on 1.8 million transcript–context profiles from 73 cellular contexts across three species, it unifies translation dynamics, TE, ORF identification, cultured-cell design and neoantigen prioritization (Supplementary Table 3).",
    8: "Raw Ribo-seq signal is sparse, protocol dependent and confounded by transcript abundance.",
    9: "We curated 1,355 Ribo-seq runs and 475 matched RNA-seq runs from 73 cellular contexts across human, rhesus macaque and mouse.",
    10: "A unified workflow performed quality control, within-context merging, representative-transcript selection, nuclease-aware P-site inference, profile filtering and RNA-abundance normalization.",
    11: "Within-context merging substantially increased per-transcript coverage, with CDS coverage of housekeeping genes approaching 1.0.",
    12: "Compositional regression reduced RNA-abundance confounding to yield relative ribosome occupancy profiles.",
    13: "Cross-species and cross-context patterns are consistent with a substantial cis-sequence contribution without establishing causal dominance. This supports an RNA-sequence backbone with cellular context conditioning.",
    14: "Processed profiles were enriched across annotated CDSs and showed strong three-nucleotide periodicity.",
    15: "The final resource contains 1.8 million transcript–context profiles, with one profile for each observed transcript and cellular context (Fig. 1).",
    16: "The resource includes protein-coding transcripts and noncoding RNAs with measurable occupancy but does not represent every translatome state.",
    18: "TRACE uses full-length RNA sequence as a shared backbone; bidirectional attention, Rotary Position Embedding, Flash Attention and length-aware batching support variable-length transcripts.",
    19: "Adaptive Layer Normalization conditions the shared sequence backbone on cellular context.",
    20: "Training combines position-level occupancy, CDS frame-aware TE and batchwise transcript-ranking objectives. CDS and frame labels supervise training, whereas inference requires no annotations.",
    21: "On chromosome-held-out transcripts, TRACE produced periodic, CDS-enriched occupancy profiles, indicating generalization of supervised translation patterns (Fig. 2b,c).",
    22: "Predicted periodicity differed between housekeeping transcripts and noncoding RNAs (Fig. 2d).",
    23: "For ENST00000332859, predicted and observed profiles reached Spearman ρ = 0.80. A localized noncoding-RNA prediction overlapped Ribo-seq evidence without establishing stable protein production.",
    24: "The same profile supports TE estimation and forward design within known CDSs and annotation-free ORF discovery across unannotated transcripts.",
    26: "TRACE was benchmarked on translation dynamics, TE and ORF identification. Among the models in Supplementary Table 3, only TRACE combines nucleotide-resolution occupancy, cellular context conditioning, unified multi-species modeling and annotation-free inference.",
    27: "For translation dynamics, TRACE approached cross-study experimental baselines and outperformed retrained sequence-only Riboformer and RiboMIMO on the evaluated datasets (Fig. 3b–d).",
    28: "For TE, TRACE ranked among the top-performing models across polysome profiling, protein-to-RNA ratios and SILAC synthesis rates (Fig. 3e–g).",
    29: "For ORF identification, TRACE achieved higher F1 than the evaluated Ribo-seq-based callers and sequence baselines without sample-specific Ribo-seq input.",
    30: "In matched fetal-brain data, TRACE profiles correlated with Ribo-seq, and short-ORF precision was comparable to RiboTIE.",
    32: "Prediction sensitivity and attention were analyzed across transcript architecture, initiation motifs and codon positions. Because CDS and frame labels supervised training, attention is interpreted associatively rather than causally.",
    33: "Predicted TE was associated with 5′ UTR, CDS and 3′ UTR length and with RNA minimum free energy; these correlations do not establish causality.",
    34: "Attention was enriched near start, CDS and stop regions and was associated with Kozak context.",
    35: "Frame-resolved attention reflected training supervision, whereas in silico codon substitutions measured prediction sensitivity rather than elongation mechanisms.",
    36: "In silico upstream-start perturbations produced prediction changes consistent with known initiation-associated patterns.",
    38: "Ten Fluc and mCherry designs showed graded cultured-cell protein output associated with TRACE predictions (Pearson r > 0.9).",
    39: "A TRACE-designed Gluc CDS produced higher expression than wild-type, RiboDecode and LinearDesign sequences in the evaluated cultured-cell experiment.",
    40: "Cellular-context-specific designs produced the intended relative expression patterns in HEK293T, HeLa and HepG2 cells.",
    42: "Non-canonical translation and alternative splicing expand candidate epitopes beyond somatic mutations but require prioritization of translated ORFs.",
    43: "Annotation-free occupancy prediction supports candidate nomination from tumor and normal RNA-seq without matched tumor Ribo-seq or mass spectrometry.",
    44: "The pipeline combined isoform assembly, tumor-specific transcript selection, TRACE ORF prediction and HLA binding prediction to nominate ten neoepitopes.",
    45: "IFN-γ ELISpot and organoid co-culture supported selected candidates functionally but did not directly quantify endogenous peptide presentation.",
    47: "TRACE combines a harmonized resource of 1.8 million transcript–context profiles with full-length sequence modeling and cellular context conditioning.",
    48: "Cross-species patterns support sequence-centered modeling, but neither these associations nor supervised attention establish causal mechanisms.",
    49: "One translation representation links forward CDS design with annotation-free reverse discovery of candidate ORFs and neoantigens.",
    50: "Limitations include reduced accuracy for low-occupancy transcripts and underrepresented contexts, the distinction between occupancy and stable protein output, and the absence of direct immunopeptidomic confirmation.",
    51: "The harmonized data representation and joint ORF and TE objectives could provide translation-focused fine-tuning for RNA foundation models.",
    52: "TRACE could guide cell-type-specific expression optimization for protein-replacement mRNAs and Cas proteins, but these applications require in vivo validation.",
    53: "Population-scale tumor-transcriptome analysis could prioritize recurrent neoantigen candidates across patients. Validation across HLA backgrounds may enable more broadly deployable mRNA vaccines and reduce, rather than eliminate, individualized design costs.",
}


NUMBER_GROUPS = [
    range(3, 7),
    range(8, 17),
    range(18, 25),
    range(26, 31),
    range(32, 37),
    range(38, 41),
    range(42, 46),
    range(47, 54),
]


def set_text_preserving_first_run(paragraph, text):
    runs = paragraph.runs
    if runs:
        first = runs[0]
        for run in runs[1:]:
            paragraph._p.remove(run._r)
        first.text = text
    else:
        paragraph.add_run(text)


def next_numeric_id(elements, attribute):
    values = [int(element.get(qn(attribute))) for element in elements]
    return max(values, default=0) + 1


def add_decimal_numbering(document):
    numbering = document.part.numbering_part.element
    abstract_id = next_numeric_id(numbering.findall(qn("w:abstractNum")), "w:abstractNumId")
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num_id = next_numeric_id(numbering.findall(qn("w:num")), "w:numId")
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:numPr"))
    if existing is not None:
        p_pr.remove(existing)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_element)
    p_pr.append(num_pr)


def clone_paragraph_after(paragraph, text):
    cloned = deepcopy(paragraph._p)
    paragraph._p.addnext(cloned)
    new_paragraph = Paragraph(cloned, paragraph._parent)
    set_text_preserving_first_run(new_paragraph, text)
    return new_paragraph


def revise_outline():
    document = Document(WORKDIR / "TRACE_outline.source.docx")
    for index, text in OUTLINE_TEXT.items():
        set_text_preserving_first_run(document.paragraphs[index], text)
    for group in NUMBER_GROUPS:
        num_id = add_decimal_numbering(document)
        for index in group:
            apply_numbering(document.paragraphs[index], num_id)
    document.save(WORKDIR / "TRACE_outline.revised.docx")


def revise_manuscript():
    document = Document(WORKDIR / "TRACE_manuscript.source.docx")
    set_text_preserving_first_run(
        document.paragraphs[56],
        "We collected 1,355 Ribo-seq runs across human, rhesus macaque and mouse, together with 475 matched RNA-seq runs. All input data were obtained as raw FASTQ files. For Ribo-seq, adapters were first detected and removed with fastp (version 0.23.4). When automatic detection was unsuitable, adapter sequences were obtained from the source publication and removed with cutadapt (version 4.5) or Trimmomatic (version 0.39). To reduce rRNA and tRNA contamination, trimmed reads were aligned with Bowtie2 (version 2.5.2) to species-specific rRNA and tRNA reference sequences (human, hg38, Ensembl V114; rhesus macaque, Ensembl V101, rheMac10; mouse, mm39, Ensembl V109). Unmapped reads were retained and aligned to the corresponding reference genome with STAR (version 2.7.9a). We retained Ribo-seq datasets for which more than 60% of aligned reads mapped within annotated coding sequences (CDSs).",
    )
    set_text_preserving_first_run(
        document.paragraphs[15],
        "All data passed through a common workflow. After run-level quality control, BAM files from the same cellular context were merged. Representative transcripts were selected from alignment support; exon bins with at least two reads were considered informative, MANE Select transcripts were retained, and non-MANE isoforms adding no informative exon or splice-junction evidence were removed. A nuclease-aware LightGBM classifier then assigned read-specific P-sites across library conditions. Profiles were subsequently filtered by RNA abundance, P-site depth, coverage and, for annotated CDSs, frame-0 enrichment. Dataset-level periodicity quality control retained libraries with clear three-nucleotide signal. Detailed procedures are provided in Supplementary Methods.",
    )
    set_text_preserving_first_run(
        document.paragraphs[51],
        "TRACE predicts ribosome occupancy rather than downstream protein stability, folding, modification or degradation. Accuracy remains limited for transcripts with very low occupancy and for cellular contexts underrepresented in training. For neoantigen candidates, functional assays do not replace direct evidence of endogenous peptide presentation.",
    )
    foundation = (
        "The harmonized data representation and joint training objectives may also provide a translation-focused adaptation framework for RNA foundation models. Fine-tuning such models with nucleotide-resolution occupancy, ORF supervision and transcript-level TE objectives could support both ORF identification and TE quantification."
    )
    therapy = (
        "TRACE could further provide a sequence-design objective for cell-type-specific protein expression, including protein-replacement mRNAs and Cas proteins used in genome editing. These applications require in vivo validation, where delivery, RNA stability, immune activation and tissue regulation may alter therapeutic protein output."
    )
    population = (
        "Finally, applying annotation-free TRACE predictions to population-scale tumor transcriptomes could prioritize recurrent neoantigen candidates shared across patients. Validation by immunopeptidomics and functional assays across diverse HLA backgrounds may support more broadly deployable mRNA vaccines and reduce, rather than eliminate, the cost and turnaround time of individualized antigen discovery and vaccine design."
    )
    set_text_preserving_first_run(document.paragraphs[53], foundation)
    current = clone_paragraph_after(document.paragraphs[53], therapy)
    clone_paragraph_after(current, population)

    context_paragraph = next(p for p in document.paragraphs if p.text.startswith("Cellular-context conditioning was further evaluated"))
    set_text_preserving_first_run(
        context_paragraph,
        "Cellular-context conditioning was further evaluated with three strategies: Zero, in which the expression vector was set to zero; Real, in which the measured expression vector was supplied; and Augment, in which expression vectors were stochastically masked or continuously interpolated toward zero during training. To assess how training-set diversity affected generalization, each strategy was trained using 5, 22 or 40 cellular contexts and evaluated on the same chromosome-held-out transcripts from 26 unseen human cellular contexts. Selected results for the five-context model ablation are summarized in Supplementary Table 2. Full model configurations, augmentation parameters and evaluation procedures are provided in Supplementary Methods.",
    )
    document.save(WORKDIR / "TRACE_manuscript.revised.docx")


def revise_supplement():
    document = Document(WORKDIR / "supplementary_information.source.docx")
    set_text_preserving_first_run(
        document.paragraphs[19],
        "Each strategy was trained on matched sets of 5, 22 or 40 human cellular contexts; the 40-context set combined 22 tissues and 18 common cell lines. All nine combinations were evaluated on the same chromosome-held-out transcripts from 26 uncommon cell lines absent from training. Generalization was summarized by nucleotide-profile Spearman ρ, periodicity-related performance, CDS-mean signal Spearman ρ and CDS-mean absolute error, weighting cellular contexts equally.",
    )
    set_text_preserving_first_run(
        document.paragraphs[20],
        "Supplementary Table 2 | Selected model-ablation performance on the five-context validation dataset",
    )
    set_text_preserving_first_run(
        document.paragraphs[21],
        "Table note. Values are the best values recorded across 50 epochs, with the corresponding epoch in parentheses. All metrics used the human_5c_6k_depth0.1_cov0.1_rpm1 validation set. Lower values are better for validation loss and CDS-mean MAE, whereas higher values are better for Spearman ρ. The best value in each column is shaded green. Because each metric was selected independently, values within a row can correspond to different checkpoints. Model names were standardized to the manuscript terminology: TRACE-Augment corresponds to TRACE (Mask+Interp.) in the source file, LN-model to LN Transformer and Conv-model to Conv model.",
    )

    table = document.tables[0]
    row_values = [
        ["TRACE-Zero", "0.1093 (29)", "0.437 (44)", "0.612 (29)", "0.0981 (26)"],
        ["TRACE-Augment", "0.1077 (41)", "0.444 (40)", "0.635 (42)", "0.0956 (41)"],
        ["LN-model", "0.1109 (42)", "0.427 (10)", "0.600 (44)", "0.1009 (11)"],
        ["Conv-model", "0.1190 (5)", "0.399 (5)", "0.587 (40)", "0.1016 (38)"],
    ]
    for row, values in zip(table.rows[1:], row_values):
        for cell, value in zip(row.cells, values):
            set_text_preserving_first_run(cell.paragraphs[0], value)
    document.save(WORKDIR / "supplementary_information.revised.docx")


def main():
    revise_outline()
    revise_manuscript()
    revise_supplement()


if __name__ == "__main__":
    main()
