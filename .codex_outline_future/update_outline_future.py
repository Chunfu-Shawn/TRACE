from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOURCE = Path(".codex_outline_future/TRACE_outline.source.docx")
OUTPUT = Path(".codex_outline_future/TRACE_outline.revised.docx")


FUTURE_PARAGRAPHS = [
    (
        "Near-term extensions include incorporating RNA structure, RNA modifications, "
        "RBP occupancy and subcellular localization, improving generalization to sparsely "
        "sampled cellular contexts, and obtaining direct immunopeptidomic validation of "
        "candidate neoantigens."
    ),
    (
        "More broadly, the harmonized data representation and joint training strategy "
        "developed here may provide a translation-focused adaptation framework for RNA "
        "foundation models. Fine-tuning such models with nucleotide-resolution ribosome "
        "occupancy profiles, ORF supervision and transcript-level TE objectives could enable "
        "a shared pretrained RNA representation to support both ORF identification and TE "
        "quantification."
    ),
    (
        "TRACE may also provide a sequence-design objective for cell-type-specific protein "
        "expression. Potential applications include optimizing mRNAs for protein-replacement "
        "therapies and tuning the expression of Cas and other genome-editing proteins. These "
        "applications remain to be tested in vivo, where delivery, RNA stability, innate immune "
        "activation and tissue-specific regulation may alter the relationship between predicted "
        "ribosome occupancy and therapeutic protein output."
    ),
    (
        "Finally, applying annotation-free TRACE predictions to population-scale tumor "
        "transcriptomes could help identify recurrent, tumor-associated translated ORFs and "
        "candidate neoepitopes shared across patients or population groups. If validated by "
        "immunopeptidomics and functional immune assays across diverse HLA backgrounds, such "
        "candidates could support more broadly deployable mRNA vaccines and reduce, rather than "
        "eliminate, the cost and turnaround time of fully individualized antigen discovery and "
        "vaccine design."
    ),
]


def set_paragraph_text_preserving_first_run(paragraph, text):
    runs = paragraph.runs
    if runs:
        first = runs[0]
        for run in runs[1:]:
            paragraph._p.remove(run._r)
        first.text = text
    else:
        paragraph.add_run(text)


def clone_paragraph_after(paragraph, text):
    cloned = deepcopy(paragraph._p)
    paragraph._p.addnext(cloned)
    new_paragraph = paragraph._parent.add_paragraph()
    placeholder = new_paragraph._p
    cloned.addnext(placeholder)
    placeholder.getparent().remove(placeholder)

    # Resolve the cloned XML element through the document paragraph collection.
    for candidate in paragraph._parent.paragraphs:
        if candidate._p is cloned:
            set_paragraph_text_preserving_first_run(candidate, text)
            return candidate
    raise RuntimeError("Unable to resolve cloned paragraph")


def main():
    doc = Document(SOURCE)
    targets = [p for p in doc.paragraphs if p.text.startswith("Future work:")]
    if len(targets) != 1:
        raise RuntimeError(f"Expected exactly one Future work paragraph, found {len(targets)}")

    current = targets[0]
    set_paragraph_text_preserving_first_run(current, FUTURE_PARAGRAPHS[0])
    for text in FUTURE_PARAGRAPHS[1:]:
        current = clone_paragraph_after(current, text)

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
