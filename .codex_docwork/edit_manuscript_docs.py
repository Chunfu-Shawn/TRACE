from copy import deepcopy
from pathlib import Path

from docx import Document


WORKDIR = Path(__file__).resolve().parent
MAIN_SOURCE = WORKDIR / "TRACE_manuscript.source.docx"
SUPP_SOURCE = WORKDIR / "supplementary_information.source.docx"
MAIN_OUTPUT = WORKDIR / "TRACE_manuscript.revised.docx"
SUPP_OUTPUT = WORKDIR / "supplementary_information.revised.docx"


def paragraph_by_text(document, exact_text):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip() == exact_text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph matching {exact_text!r}, found {len(matches)}")
    return matches[0]


def paragraph_starting_with(document, prefix):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting with {prefix!r}, found {len(matches)}")
    return matches[0]


def replace_paragraph_text(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def formula_text(paragraph):
    return "".join(paragraph._p.xpath(".//m:t/text()"))


main = Document(MAIN_SOURCE)
supp = Document(SUPP_SOURCE)

# Expand the main normalization method only where the scientific principle is essential.
selection_paragraph = paragraph_starting_with(main, "For each cellular context, Ribo-seq alignment support")
replace_paragraph_text(
    selection_paragraph,
    "For each cellular context, Ribo-seq alignment support was used to select a parsimonious set of representative transcripts that captured supported exons and splice junctions, while retaining MANE Select isoforms. Transcripts that did not explain additional informative alignment features were removed, reducing isoform redundancy before profile construction.",
)

normalization_summary = paragraph_starting_with(main, "Finally, P-site profiles were normalized")
replace_paragraph_text(
    normalization_summary,
    "We then used a nuclease-aware LightGBM classifier to infer a read-specific P-site offset from fragment length and sequence context at both read ends. This formulation accommodated libraries generated under different nuclease conditions, including RNase I and micrococcal nuclease (MNase), without imposing a single length-specific offset rule. P-site assignments were aggregated at nucleotide resolution. Informative transcript–context profiles were retained using matched RNA-seq abundance and P-site depth and coverage; frame periodicity was additionally required only for transcripts with a valid annotated CDS, whereas full-length criteria were used otherwise.",
)

architecture_heading = paragraph_by_text(main, "TRACE model architecture")
architecture_heading.insert_paragraph_before(
    "To construct the model target, non-zero P-site counts were winsorized at the 99.5th percentile. Within each cellular context, transcript-level RPF and matched RNA-seq counts were augmented by a pseudocount of 10 and centered-log-ratio (CLR) transformed. An ordinary least-squares regression of CLR-transformed RPF counts on CLR-transformed RNA-seq counts yielded residual rₜ, and exp(rₜ) provided a multiplicative transcript-level scale for translation not explained by measured RNA abundance. For transcript t and nucleotide position ℓ, the RNA-abundance-adjusted relative ribosome occupancy target was defined as:",
    style="Normal",
)

supp_formula_paragraphs = [paragraph for paragraph in supp.paragraphs if paragraph.style.name == "公式"]
if len(supp_formula_paragraphs) < 2:
    raise RuntimeError("The supplementary document does not contain the expected normalization formulas")
for formula_paragraph in supp_formula_paragraphs[:2]:
    architecture_heading._p.addprevious(deepcopy(formula_paragraph._p))

architecture_heading.insert_paragraph_before(
    "Here c⁽ʷ⁾ₜₗ denotes the winsorized P-site count, Pₜ⁺ is the set of positions with a positive winsorized count, c̄ₜ⁽⁺⁾ is the mean over those positions, Rₜ is the transcript-level RPF count and Aₜ is the matched RNA-seq count. The regression coefficients β₀ and β₁ were estimated separately within each cellular context. Thus, yₜₗ is a dimensionless signal that combines within-transcript positional occupancy with an RNA-abundance-adjusted transcript-level scale; it is not an absolute ribosome density, a single-molecule measurement or a site-occupancy probability.",
    style="Normal",
)
architecture_heading.insert_paragraph_before(
    "This workflow yielded 1.8 million transcript–context profiles from 73 cellular contexts across human, rhesus macaque and mouse. A transcript observed in multiple cellular contexts contributed one profile per context. Detailed transcript selection, P-site inference, profile filtering and normalization parameters are provided in Supplementary Methods.",
    style="Normal",
)

# Add a concise main-text ablation design after model training.
benchmark_heading = paragraph_by_text(main, "Benchmarking on translation tasks")
benchmark_heading.insert_paragraph_before("Ablation experiments", style="Heading 2")
benchmark_heading.insert_paragraph_before(
    "We used matched-input ablations to distinguish the contributions of cellular-context conditioning and sequence modeling. BaseModelLN was a sequence-only pre-LayerNorm Transformer that retained global self-attention but removed adaptive cellular-context conditioning. BaseModelConv replaced the Transformer encoder with parameter-matched residual one-dimensional convolutional blocks, providing a local-sequence baseline. All variants used the same nucleotide input, prediction head, training objectives and chromosome-disjoint evaluation framework. Comparisons among TRACE-Real, TRACE-Zero and BaseModelLN distinguished measured cellular-context information from the adaptive-normalization parameterization, whereas the BaseModelLN–BaseModelConv comparison tested global self-attention against local convolution.",
    style="Normal",
)

# Correct an adjacent punctuation typo without changing the surrounding method.
for paragraph in main.paragraphs:
    if "targets.We used the same human chromosome" in paragraph.text:
        replace_paragraph_text(
            paragraph,
            paragraph.text.replace("targets.We used", "targets. We used"),
        )
        break
benchmark_heading.insert_paragraph_before(
    "Cellular-context conditioning was further evaluated with three strategies: Zero, in which the expression vector was set to zero; Real, in which the measured expression vector was supplied; and Expression augment, in which expression vectors were stochastically masked or continuously interpolated toward zero during training. To assess how training-set diversity affected generalization, each strategy was trained using 5, 22 or 40 cellular environments and evaluated on the same chromosome-held-out transcripts from 26 unseen human cellular environments. Full model configurations, augmentation parameters and evaluation procedures are provided in Supplementary Methods.",
    style="Normal",
)

# Add implementation-level ablation details to Supplementary Methods.
table_caption = paragraph_starting_with(supp, "Supplementary Table 2 |")
table_caption.insert_paragraph_before("Ablation experiments", style="Heading 2")
table_caption.insert_paragraph_before(
    "Structural ablations retained TRACE's nucleotide embedding, prediction head, losses, chromosome partitions and optimization. BaseModelLN used 12 standard pre-LayerNorm Transformer layers (model dimension, 384; attention heads, 16; feed-forward dimension, 768; dropout, 0.1) with sequence and padding-mask inputs only. BaseModelConv used 12 residual one-dimensional convolutional blocks (pre-LayerNorm; kernel width, 7; hidden dimension, 384; GELU; 1 × 1 projection; dropout, 0.1). Its encoder parameter count was within 0.3% of BaseModelLN. Both models ignored cellular-context and species inputs.",
    style="Normal",
)
table_caption.insert_paragraph_before(
    "The cellular-context ablation used the TRACE backbone. Zero forced expression to zero during training and validation; Real supplied the measured 16,840-gene vector without perturbation. In Expression augment, 10% of training samples received an exact zero vector; among the remainder, 30% were scaled by a value sampled uniformly from 0 to 1 after adding Gaussian noise (s.d., 0.15). Real versus Zero tested measured context, TRACE-Zero versus BaseModelLN tested adaptive normalization without measured expression, and BaseModelLN versus BaseModelConv tested global self-attention against local convolution.",
    style="Normal",
)
last_ablation_paragraph = table_caption.insert_paragraph_before(
    "Each strategy was trained on matched sets of 5, 22 or 40 human cellular environments; the 40-environment set combined 22 tissues and 18 common cell lines. All nine combinations were evaluated on the same chromosome-held-out transcripts from 26 uncommon cell lines absent from training. Generalization was summarized by nucleotide-profile Spearman ρ, periodicity-related performance, CDS-mean signal Spearman ρ and CDS-mean absolute error, weighting cellular environments equally.",
    style="Normal",
)

# Keep Supplementary Methods portrait and preserve the comparison table's landscape section.
normalization_end = paragraph_starting_with(supp, "Here c⁽ʷ⁾ₜₗ denotes the winsorized P-site count")
source_ppr = normalization_end._p.get_or_add_pPr()
section_properties = source_ppr.sectPr
if section_properties is None:
    raise RuntimeError("Expected a section break before Supplementary Table 2")
source_ppr.remove(section_properties)
target_ppr = last_ablation_paragraph._p.get_or_add_pPr()
target_ppr.append(section_properties)

main.save(MAIN_OUTPUT)
supp.save(SUPP_OUTPUT)

# Structural checks for the edited artifacts.
main_check = Document(MAIN_OUTPUT)
supp_check = Document(SUPP_OUTPUT)
assert sum(p.text.strip() == "Ablation experiments" for p in main_check.paragraphs) == 1
assert sum(p.text.strip() == "Ablation experiments" for p in supp_check.paragraphs) == 1
main_formula_texts = [formula_text(p) for p in main_check.paragraphs if p._p.xpath(".//m:oMath | .//m:oMathPara")]
assert any("exp" in text and "clr" not in text for text in main_formula_texts)
assert any("clr" in text for text in main_formula_texts)
assert "26 unseen human cellular environments" in " ".join(p.text for p in main_check.paragraphs)
assert "1.8 million transcript–context profiles" in " ".join(p.text for p in main_check.paragraphs)
assert len(main_check.tables) == 0
assert len(supp_check.tables) == 1

print(MAIN_OUTPUT)
print(SUPP_OUTPUT)
